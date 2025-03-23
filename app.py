import os
import io
import json
import datetime
import streamlit as st
from typing import Dict, Any, List, Optional
import openai
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64

# Load environment variables
load_dotenv()

# Set OpenAI API key
openai.api_key = os.getenv("OPENAI_API_KEY")

class LOAGenerator:
    """
    A class for generating Letters of Authorization (LOAs) for outdoor advertising
    based on specific parameters using OpenAI's API.
    """
    
    def __init__(self):
        self.model = "gpt-4"  # Can be changed to gpt-3.5-turbo for lower cost
        self.conversation_history = []
        self.current_loa = None
        
    def _create_system_prompt(self) -> str:
        """
        Creates the system prompt that instructs the model on how to generate LOAs.
        
        Returns:
            str: The system prompt
        """
        return """
        You are an expert in creating Letters of Authorization (LOAs) for outdoor advertising agencies. 
        Your task is to generate professional, legally-sound LOAs based on the parameters provided.
        
        Follow these guidelines:
        1. Use formal business language appropriate for official documents
        2. Structure the LOA with proper sections including reference numbers, dates, recipient details, subject line, main body, and signatory information
        3. Include all necessary legal clauses regarding installation, maintenance, payments, and liability
        4. Format dates as DD.MM.YYYY
        5. Make the content specific to the scenario provided
        6. Ensure payment terms and conditions are clearly stated
        7. Include appropriate references to any tenders or previous communications when provided
        
        Return only the plain text content of the LOA without any explanations or additional formatting.
        """
    
    def _construct_loa_prompt(self, params: Dict[str, Any]) -> str:
        """
        Constructs a prompt for the model to generate an LOA based on the provided parameters.
        
        Args:
            params: Dictionary of parameters for LOA generation
            
        Returns:
            str: The constructed prompt
        """
        # Format date
        date_str = datetime.date.today().strftime("%d.%m.%Y")
        
        # Construct the prompt
        prompt = f"""
        Generate a Letter of Authorization (LOA) with the following details:
        
        Date: {date_str}
        
        Recipient Address:
        {params.get('address', '')}
        
        To: {params.get('to_whom', 'To Whom It May Concern')}
        
        Scenario: {params.get('scenario', '')}
        
        Specific Details to Include:
        {params.get('specific_details', '')}
        
        Closing:
        Yours sincerely,
        {params.get('yours_sincerely', '')}
        
        The LOA should follow a formal business letter format with:
        1. Clear header with date
        2. Recipient address
        3. Appropriate salutation
        4. A clear introduction stating the purpose of the letter
        5. A main body detailing the specifics of the authorization
        6. Any necessary terms and conditions
        7. A formal closing with signatory information
        
        Ensure the content specifically addresses the scenario provided and incorporates all the specific details mentioned.
        """
        
        return prompt
    
    def generate_loa(self, params: Dict[str, Any]) -> str:
        """
        Generates an LOA based on the provided parameters.
        
        Args:
            params: Dictionary of parameters for LOA generation
            
        Returns:
            str: The generated LOA text
        """
        # Create messages for the API call
        messages = [
            {"role": "system", "content": self._create_system_prompt()},
            {"role": "user", "content": self._construct_loa_prompt(params)}
        ]
        
        # Add conversation history if available
        if self.conversation_history:
            messages.extend(self.conversation_history)
        
        try:
            # Call the OpenAI API
            response = openai.chat.completions.create(
                model=self.model,
                messages=messages,
                temperature=0.2,  # Lower temperature for more consistent outputs
                max_tokens=2500
            )
            
            # Get the generated LOA content
            loa_content = response.choices[0].message.content
            
            # Store the conversation
            self.conversation_history = [
                {"role": "user", "content": self._construct_loa_prompt(params)},
                {"role": "assistant", "content": loa_content}
            ]
            
            # Store the current LOA
            self.current_loa = loa_content
            
            return loa_content
        
        except Exception as e:
            return f"Error generating LOA: {str(e)}"
    
    def edit_loa(self, edit_request: str) -> str:
        """
        Edits the previously generated LOA based on the edit request.
        
        Args:
            edit_request: Description of the edits to make
            
        Returns:
            str: The edited LOA text
        """
        if not self.current_loa:
            return "No LOA has been generated yet. Please generate an LOA first."
        
        # Create the edit request message
        edit_message = {
            "role": "user", 
            "content": f"""
            Edit the LOA according to the following request:
            
            {edit_request}
            
            Return the complete edited LOA.
            """
        }
        
        # Add the edit request to the conversation history
        self.conversation_history.append(edit_message)
        
        # Create messages for the API call
        messages = [
            {"role": "system", "content": self._create_system_prompt()},
        ]
        
        # Add conversation history
        messages.extend(self.conversation_history)
        
        try:
            # Call the OpenAI API
            response = openai.chat.completions.create(
                model=self.model,
                messages=messages,
                temperature=0.2,
                max_tokens=2500
            )
            
            # Get the edited LOA content
            edited_loa = response.choices[0].message.content
            
            # Update the conversation history
            self.conversation_history.append({"role": "assistant", "content": edited_loa})
            
            # Update the current LOA
            self.current_loa = edited_loa
            
            return edited_loa
        
        except Exception as e:
            return f"Error editing LOA: {str(e)}"


def create_word_document(loa_content: str) -> Document:
    """
    Create a Word document from the LOA content.
    
    Args:
        loa_content: The text content of the LOA
        
    Returns:
        Document: A python-docx Document object containing the formatted LOA
    """
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Split the content into lines
    lines = loa_content.strip().split('\n')
    
    # Process each line
    for line in lines:
        line = line.strip()
        if not line:
            # Add empty paragraph for blank lines
            doc.add_paragraph()
        elif line.startswith('Date:') or line.lower().startswith('ref:') or line.lower().startswith('reference:'):
            # Right-aligned date/reference
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(line).bold = True
        elif any(line.lower().startswith(salutation) for salutation in ['dear ', 'to whom', 'sir', 'madam']):
            # Salutation
            p = doc.add_paragraph()
            p.add_run(line)
        elif line.lower().startswith('subject:'):
            # Subject line - bold
            p = doc.add_paragraph()
            p.add_run(line).bold = True
        elif line.lower().startswith(('yours ', 'sincerely', 'faithfully', 'regards')):
            # Closing
            p = doc.add_paragraph()
            p.add_run(line)
        elif line.endswith(':') and len(line) < 50:
            # Section headers
            p = doc.add_paragraph()
            p.add_run(line).bold = True
        else:
            # Regular paragraph
            p = doc.add_paragraph()
            p.add_run(line)
    
    return doc


def get_docx_download_link(doc, filename="letter_of_authorization.docx"):
    """
    Generate a download link for a Word document.
    
    Args:
        doc: The python-docx Document object
        filename: The name for the downloaded file
        
    Returns:
        str: HTML for the download link
    """
    # Save the document to a BytesIO object
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # Get b64 encoded string
    b64 = base64.b64encode(file_stream.read()).decode()
    
    # Generate the download link
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Word Document</a>'
    
    return href


def main():
    st.set_page_config(page_title="LOA Generator", page_icon="📝", layout="wide")
    
    # Add custom CSS
    st.markdown("""
        <style>
        .main .block-container {
            padding-top: 2rem;
        }
        .stTextArea textarea {
            height: 150px;
        }
        .loa-preview {
            border: 1px solid #ddd;
            border-radius: 5px;
            padding: 20px;
            background-color: #f9f9f9;
            font-family: 'Times New Roman', Times, serif;
            white-space: pre-line;
            margin-bottom: 20px;
        }
        h1 {
            color: #2c3e50;
        }
        .stButton button {
            background-color: #2c3e50;
            color: white;
        }
        </style>
        """, unsafe_allow_html=True)
    
    st.title("Letter of Authorization (LOA) Generator")
    st.subheader("Generate professional LOAs for outdoor advertising scenarios")
    
    # Initialize session state
    if 'loa_generator' not in st.session_state:
        st.session_state.loa_generator = LOAGenerator()
    
    if 'current_loa' not in st.session_state:
        st.session_state.current_loa = None
    
    if 'edit_mode' not in st.session_state:
        st.session_state.edit_mode = False
    
    # Create two columns for input form and LOA preview
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("LOA Information")
        
        # Recipient Address
        st.subheader("Recipient Address")
        address = st.text_area("Enter the complete address", height=100,
                              placeholder="Company Name\nStreet Address\nCity, State ZIP/Postal Code")
        
        # To Whom
        st.subheader("Addressee")
        to_whom = st.text_input("To whom is this addressed?", 
                               placeholder="To Whom It May Concern / Mr. John Doe / The Manager")
        
        # Scenario
        st.subheader("Scenario")
        scenario = st.text_area("Describe the advertising scenario", height=100,
                              placeholder="Example: Digital hoarding at a retail outlet for 5 years with quarterly payments")
        
        # Specific Details
        st.subheader("Specific Details")
        specific_details = st.text_area("Enter specific details to include in the LOA", height=150,
                                      placeholder="Example: Size: 20' X 20' = 400 Sq. Ft.\nPayment: Rs. 310 per sq ft per year\nLocation: FC Road, Pune\nAny other terms or conditions...")
        
        # Yours Sincerely
        st.subheader("Signature")
        yours_sincerely = st.text_area("Enter the signatory information", height=80,
                                     placeholder="Name\nPosition\nCompany Name")
        
        # Generate Button
        if st.button("Generate LOA"):
            with st.spinner("Generating LOA..."):
                # Prepare parameters
                params = {
                    "address": address,
                    "to_whom": to_whom,
                    "scenario": scenario,
                    "specific_details": specific_details,
                    "yours_sincerely": yours_sincerely
                }
                
                # Generate LOA
                loa_content = st.session_state.loa_generator.generate_loa(params)
                st.session_state.current_loa = loa_content
                st.session_state.edit_mode = False
    
    with col2:
        st.header("LOA Preview")
        
        if st.session_state.current_loa:
            # Display LOA in a styled container
            st.markdown('<div class="loa-preview">' + st.session_state.current_loa.replace('\n', '<br>') + '</div>', unsafe_allow_html=True)
            
            # Create Word document
            doc = create_word_document(st.session_state.current_loa)
            
            # Provide download link
            st.markdown(get_docx_download_link(doc), unsafe_allow_html=True)
            
            # Edit option
            if not st.session_state.edit_mode:
                if st.button("Edit LOA"):
                    st.session_state.edit_mode = True
            else:
                edit_request = st.text_area("Describe what changes you want to make", height=100,
                                          placeholder="Example: Please change the payment amount to Rs. 350 and add a clause about security deposit.")
                
                col1, col2 = st.columns([1, 1])
                with col1:
                    if st.button("Apply Changes"):
                        with st.spinner("Editing LOA..."):
                            edited_loa = st.session_state.loa_generator.edit_loa(edit_request)
                            st.session_state.current_loa = edited_loa
                            st.experimental_rerun()
                
                with col2:
                    if st.button("Cancel"):
                        st.session_state.edit_mode = False
                        st.experimental_rerun()
        else:
            st.info("Fill out the form and click 'Generate LOA' to create a new Letter of Authorization.")


if __name__ == "__main__":
    main()